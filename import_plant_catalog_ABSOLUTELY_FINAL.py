#!/usr/bin/env python3
"""
Import rostlinných dat z DOCX ceníku do databáze plant_catalog
FINÁLNÍ OPRAVENÁ VERZE - řeší všechny problémy
"""
import sqlite3
import docx
import re
import sys

def parse_plant_name(full_name, current_genus=None):
    """
    Parsuje název rostliny na latinský název, odrůdu a velikost kontejneru
    
    Args:
        full_name: např. "A.caerulea 'Blue Star' - K9" nebo "Acanthus hungaricus - K9 - paznehtník"
        current_genus: aktuální rod (např. "Aquilegia") pro expandování zkratek
    
    Returns:
        (latin_name, variety, container)
    """
    # Najdi kontejner (může být i uprostřed před českým popisem)
    # Např: "Acanthus hungaricus - K9 - paznehtník" -> najde "K9"
    container_match = re.search(r'\s*-\s*(K\d+[a-z]?)', full_name, re.IGNORECASE)
    container = container_match.group(1) if container_match else None
    
    # Název je vše před kontejnerem
    if container_match:
        name_part = full_name[:container_match.start()].strip()
    else:
        name_part = full_name.strip()
    
    # Najdi odrůdu v uvozovkách
    variety_match = re.search(r"['\"''„""‚]([^'\"''„""‚]+)['\"''„""‚]", name_part)
    variety = variety_match.group(1) if variety_match else None
    
    # Zbytek je latinský název
    if variety_match:
        latin = name_part[:variety_match.start()].strip()
    else:
        latin = name_part.strip()
    
    # Zkontroluj, jestli začíná zkratkou (např. "A." nebo "A.(C.)")
    abbr_match = re.match(r'^([A-Z]\.(?:\([A-Z]\.\))?)\s*', latin)
    
    if abbr_match and current_genus:
        # Je to zkratka - expanduj pomocí aktuálního rodu
        abbr = abbr_match.group(1)
        rest = latin[len(abbr):].strip()
        latin = current_genus + " " + rest
    
    return latin.strip(), variety, container


def is_genus_header(name):
    """
    Zjisti, zda je řádek nadpis rodu (ne konkrétní rostlina)
    """
    # Pokud má kontejner, není to nadpis
    if re.search(r'-\s*K\d+', name):
        return False
    
    # Pokud je to uppercase název bez dalšího, je to kategorie
    if name.isupper():
        return True
    
    # Pokud obsahuje pomlčku a český popis (lowercase), je to nadpis rodu
    if re.search(r'\s*-\s*[a-záčďéěíňóřšťúůýž]', name):
        return True
    
    return False


def extract_genus_from_header(name):
    """
    Extrahuj rod z nadpisu
    "Acaena -plazilka" -> "Acaena"
    "Aconitum - oměj" -> "Aconitum"
    """
    match = re.match(r'^([A-Z][a-z]+)', name)
    if match:
        return match.group(1)
    return None


def import_from_docx(docx_path, db_path):
    """Importuje data z DOCX do SQLite databáze"""
    
    print(f"Načítám dokument: {docx_path}")
    doc = docx.Document(docx_path)
    
    print(f"Připojuji se k databázi: {db_path}")
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    imported = 0
    skipped = 0
    errors = 0
    current_genus = None
    
    print(f"Zpracovávám {len(doc.tables)} tabulek...")
    
    for table_idx, table in enumerate(doc.tables):
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        if 'Název (kontejner)' not in headers:
            continue
        
        for row_idx, row in enumerate(table.rows[1:], start=1):
            cells = [cell.text.strip() for cell in row.cells]
            
            if len(cells) < 11:
                continue
            
            name_full = cells[0]
            
            if not name_full or name_full == '':
                continue
            
            # Zkontroluj, zda je to nadpis rodu
            if is_genus_header(name_full):
                genus = extract_genus_from_header(name_full)
                if genus:
                    current_genus = genus
                    print(f"  → Nový rod: {current_genus}")
                continue
            
            # Přeskoč řádky bez kontejneru
            if not re.search(r'-\s*K\d+', name_full):
                continue
            
            # Přeskoč řádky bez ceny
            if not cells[1] or cells[1] == '':
                continue
            
            try:
                latin_name, variety, container = parse_plant_name(name_full, current_genus)
                
                # Debug - ukázat první pár importů
                if imported < 3:
                    print(f"  Debug: '{name_full}' → latin='{latin_name}', variety='{variety}', container='{container}'")
                
                data = {
                    'latin_name': latin_name,
                    'variety': variety,
                    'container_size': container,
                    'flower_color': cells[2] if len(cells) > 2 and cells[2] else None,
                    'flowering_time': cells[3] if len(cells) > 3 and cells[3] else None,
                    'leaf_color': cells[4] if len(cells) > 4 and cells[4] else None,
                    'height': cells[5] if len(cells) > 5 and cells[5] else None,
                    'light_requirements': cells[6] if len(cells) > 6 and cells[6] else None,
                    'site_type': cells[7] if len(cells) > 7 and cells[7] else None,
                    'plants_per_m2': cells[8] if len(cells) > 8 and cells[8] else None,
                    'hardiness_zone': cells[9] if len(cells) > 9 and cells[9] else None,
                    'notes': cells[10] if len(cells) > 10 and cells[10] else None
                }
                
                cursor.execute('''
                    INSERT OR IGNORE INTO plant_catalog 
                    (latin_name, variety, container_size, flower_color, flowering_time,
                     leaf_color, height, light_requirements, site_type, plants_per_m2,
                     hardiness_zone, notes)
                    VALUES (:latin_name, :variety, :container_size, :flower_color, :flowering_time,
                            :leaf_color, :height, :light_requirements, :site_type, :plants_per_m2,
                            :hardiness_zone, :notes)
                ''', data)
                
                if cursor.rowcount > 0:
                    imported += 1
                    if imported % 50 == 0:
                        print(f"  Importováno: {imported} rostlin...")
                else:
                    skipped += 1
                    
            except Exception as e:
                errors += 1
                print(f"  [CHYBA] Řádek {row_idx}: {name_full}")
                print(f"          {e}")
    
    conn.commit()
    conn.close()
    
    print(f"\n✓ Import dokončen!")
    print(f"  Importováno: {imported} rostlin")
    print(f"  Přeskočeno: {skipped} (duplicity)")
    print(f"  Chyby: {errors}")
    
    return imported, skipped, errors


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Použití: python3 import_plant_catalog.py <cesta_k_docx> <cesta_k_databazi>")
        print("Příklad: python3 import_plant_catalog.py cenik.docx app.db")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    db_path = sys.argv[2]
    
    import_from_docx(docx_path, db_path)
