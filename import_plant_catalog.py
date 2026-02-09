#!/usr/bin/env python3
"""
Import rostlinných dat z DOCX ceníku do databáze plant_catalog
OPRAVENÁ VERZE - sleduje aktuální rod místo pevných zkratek
"""
import sqlite3
import docx
import re
import sys

def parse_plant_name(full_name, current_genus=None):
    """
    Parsuje název rostliny na latinský název, odrůdu a velikost kontejneru
    
    Args:
        full_name: např. "A.caerulea 'Blue Star' - K9"
        current_genus: aktuální rod (např. "Aquilegia") pro expandování zkratek
    
    Returns:
        (latin_name, variety, container)
    """
    # Odstraň kontejner na konci
    container_match = re.search(r'\s*-\s*(K\d+[a-z]?)\s*$', full_name, re.IGNORECASE)
    container = container_match.group(1) if container_match else None
    name_part = full_name[:container_match.start()] if container_match else full_name
    
    # Najdi odrůdu v uvozovkách
    variety_match = re.search(r"['\"']([^'\"']+)['\"']", name_part)
    variety = variety_match.group(1) if variety_match else None
    
    # Zbytek je latinský název
    if variety_match:
        latin = name_part[:variety_match.start()].strip()
    else:
        latin = name_part.strip()
    
    # Zkontroluj, jestli začíná zkratkou (např. "A." nebo "A.(C.)")
    abbr_match = re.match(r'^([A-Z]\.(?:\([A-Z]\.\))?)\s*', latin)
    
    if abbr_match and current_genus:
        # Je to zkratka - expanduj pomocí aktuálního rodu
        abbr = abbr_match.group(1)
        rest = latin[len(abbr):].strip()
        latin = current_genus + " " + rest
    
    return latin.strip(), variety, container


def is_genus_header(name):
    """
    Zjisti, zda je řádek nadpis rodu (ne konkrétní rostlina)
    Příklady:
      - "Acaena -plazilka" -> True (rod + popis)
      - "Aconitum - oměj" -> True
      - "TRVALKY" -> True (kategorie)
      - "Acaena buchananii - K9" -> False (má kontejner = konkrétní rostlina)
    """
    # Pokud má kontejner, není to nadpis
    if re.search(r'-\s*K\d+', name):
        return False
    
    # Pokud je to uppercase název bez dalšího, je to kategorie
    if name.isupper():
        return True
    
    # Pokud obsahuje pomlčku a český popis (lowercase), je to nadpis rodu
    # Např: "Acaena -plazilka" nebo "Aconitum - oměj"
    if re.search(r'\s*-\s*[a-záčďéěíňóřšťúůýž]', name):
        return True
    
    return False


def extract_genus_from_header(name):
    """
    Extrahuj rod z nadpisu
    "Acaena -plazilka" -> "Acaena"
    "Aconitum - oměj" -> "Aconitum"
    """
    # Vezmi první slovo před pomlčkou nebo mezerou
    match = re.match(r'^([A-Z][a-z]+)', name)
    if match:
        return match.group(1)
    return None


def import_from_docx(docx_path, db_path):
    """Importuje data z DOCX do SQLite databáze"""
    
    # Načti dokument
    print(f"Načítám dokument: {docx_path}")
    doc = docx.Document(docx_path)
    
    # Připoj se k databázi
    print(f"Připojuji se k databázi: {db_path}")
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    imported = 0
    skipped = 0
    errors = 0
    current_genus = None  # Sleduj aktuální rod
    
    print(f"Zpracovávám {len(doc.tables)} tabulek...")
    
    for table_idx, table in enumerate(doc.tables):
        # První řádek je hlavička
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        # Kontrola, zda je to rostlinná tabulka
        if 'Název (kontejner)' not in headers:
            continue
        
        # Zpracuj jednotlivé řádky
        for row_idx, row in enumerate(table.rows[1:], start=1):
            cells = [cell.text.strip() for cell in row.cells]
            
            if len(cells) < 11:
                continue
            
            name_full = cells[0]
            
            # Přeskoč prázdné řádky
            if not name_full or name_full == '':
                continue
            
            # Zkontroluj, zda je to nadpis rodu
            if is_genus_header(name_full):
                genus = extract_genus_from_header(name_full)
                if genus:
                    current_genus = genus
                    print(f"  → Nový rod: {current_genus}")
                continue
            
            # Přeskoč řádky bez kontejneru (nejsou to konkrétní rostliny)
            if not re.search(r'-\s*K\d+', name_full):
                continue
            
            # Přeskoč řádky bez ceny (nejsou prodejní)
            if not cells[1] or cells[1] == '':
                continue
            
            try:
                latin_name, variety, container = parse_plant_name(name_full, current_genus)
                
                # Připrav data
                data = {
                    'latin_name': latin_name,
                    'variety': variety,
                    'container_size': container,
                    'flower_color': cells[2] if len(cells) > 2 else None,
                    'flowering_time': cells[3] if len(cells) > 3 else None,
                    'leaf_color': cells[4] if len(cells) > 4 else None,
                    'height': cells[5] if len(cells) > 5 else None,
                    'light_requirements': cells[6] if len(cells) > 6 else None,
                    'site_type': cells[7] if len(cells) > 7 else None,
                    'plants_per_m2': cells[8] if len(cells) > 8 else None,
                    'hardiness_zone': cells[9] if len(cells) > 9 else None,
                    'notes': cells[10] if len(cells) > 10 else None
                }
                
                # Vložit do databáze
                cursor.execute('''
                    INSERT OR IGNORE INTO plant_catalog 
                    (latin_name, variety, container_size, flower_color, flowering_time,
                     leaf_color, height, light_requirements, site_type, plants_per_m2,
                     hardiness_zone, notes)
                    VALUES (:latin_name, :variety, :container_size, :flower_color, :flowering_time,
                            :leaf_color, :height, :light_requirements, :site_type, :plants_per_m2,
                            :hardiness_zone, :notes)
                ''', data)
                
                if cursor.rowcount > 0:
                    imported += 1
                    if imported % 50 == 0:
                        print(f"  Importováno: {imported} rostlin...")
                else:
                    skipped += 1
                    
            except Exception as e:
                errors += 1
                print(f"  [CHYBA] Řádek {row_idx}: {name_full}")
                print(f"          {e}")
    
    # Commit změn
    conn.commit()
    conn.close()
    
    print(f"\n✓ Import dokončen!")
    print(f"  Importováno: {imported} rostlin")
    print(f"  Přeskočeno: {skipped} (duplicity)")
    print(f"  Chyby: {errors}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Použití: python3 import_plant_catalog.py <cesta_k_docx> <cesta_k_databazi>")
        print("Příklad: python3 import_plant_catalog.py cenik.docx app.db")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    db_path = sys.argv[2]
    
    import_from_docx(docx_path, db_path)
